import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_styles(doc):
    """Configureaza stilurile globale conform BunePractici.md"""
    # Stil Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.5
    
    # Stil Heading 1 (Capitole)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = None # Negru
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(18)
    h1._element.pPr.get_or_add_outlineLvl().val = 0

    # Stil Heading 2 (Subcapitole)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = None # Negru
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(12)
    h2._element.pPr.get_or_add_outlineLvl().val = 1

def add_toc(paragraph):
    """Insereaza codul XML pentru Table of Contents"""
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)

    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar)

    run = paragraph.add_run("Update Table (Click dreapta -> Update Field)")
    
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)

def add_formatted_text(paragraph, text):
    """Proceseaza textul pentru a aplica Bold si Italic (Markdown style)"""
    # Curatam textul de simboluri LaTeX simple pentru a nu bugui randarea textului
    text = text.replace('$$', '').replace('$', '')
    
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
            
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def build_docx():
    md_path = os.path.join('Licenta', 'Lucrare_Licenta_v2.md')
    docx_path = os.path.join('Licenta', 'Lucrare_Licenta_Finala.docx')
    
    if not os.path.exists(md_path):
        print(f"Eroare: Nu am gasit {md_path}")
        return

    doc = Document()
    setup_styles(doc)
    
    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Adaugare pagina pentru Cuprins
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CUPRINS")
    run.bold = True
    run.font.size = Pt(16)
    
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        clean_line = line.strip()
        
        if not clean_line:
            i += 1
            continue
            
        # Ignora separatoarele (---)
        if clean_line == '---':
            i += 1
            continue

        # Detectie Tabel Markdown
        if clean_line.startswith('|'):
            # Colectam toate liniile tabelului
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 0:
                # Filtram linia de separare (|---|---|)
                rows_data = []
                for tl in table_lines:
                    if re.match(r'^\|[\s:-|]+\|$', tl):
                        continue
                    cells = [c.strip() for c in tl.split('|') if c.strip() or tl.count('|') > 2]
                    # Fix empty first/last cell due to split on |
                    if tl.startswith('|'): cells = [c.strip() for c in tl[1:-1].split('|')]
                    rows_data.append(cells)
                
                if rows_data:
                    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                    table.style = 'Table Grid'
                    for r_idx, row_cells in enumerate(rows_data):
                        for c_idx, cell_text in enumerate(row_cells):
                            if c_idx < len(table.columns):
                                cell = table.cell(r_idx, c_idx)
                                cell_p = cell.paragraphs[0]
                                add_formatted_text(cell_p, cell_text)
            continue

        # Procesare titluri
        if clean_line.startswith('#'):
            level = 0
            while level < len(clean_line) and clean_line[level] == '#':
                level += 1
            
            title_text = clean_line.lstrip('#').strip()
            
            if level == 1:
                doc.add_page_break()
                p = doc.add_heading(title_text, level=1)
            elif level == 2:
                p = doc.add_heading(title_text, level=2)
            else:
                p = doc.add_heading(title_text, level=min(level, 9))
            
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            i += 1
            continue
            
        # Liste
        if clean_line.startswith('* ') or clean_line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text_to_format = re.sub(r'^[\*\-]\s+', '', clean_line)
            add_formatted_text(p, text_to_format)
            i += 1
            continue
            
        if re.match(r'^\d+\.\s', clean_line):
            # Folosim un paragraf normal pentru a pastra numerotarea originala din MD
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_formatted_text(p, clean_line)
            i += 1
            continue

        # Detectie Formule Matematice (Blocuri $$)
        if clean_line.startswith('$$') and clean_line.endswith('$$'):
            formula = clean_line.replace('$$', '').strip()
            # Curatare simboluri LaTeX comune pentru aspect vizual mai bun
            replacements = {
                '\\cdot': '·',
                '\\omega': 'ω',
                '\\frac': '',
                '{': '(',
                '}': ')',
                '\\theta': 'θ',
                '\\times': '×',
                '\\partial': '∂',
                '\\sum': 'Σ',
                '\\infty': '∞',
                '\\text': '',
                '^': ' ',
                '_': ''
            }
            for tex, char in replacements.items():
                formula = formula.replace(tex, char)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.italic = True
            run.font.name = 'Cambria Math' # Fontul standard pentru ecuatii in Word
            i += 1
            continue

        # Paragraf normal
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1)
        add_formatted_text(p, clean_line)
        i += 1

    doc.save(docx_path)
    print(f"Succes! Documentul cu Table of Contents a fost generat in: {docx_path}")

if __name__ == '__main__':
    build_docx()
